import os
import re
from collections import defaultdict

from docx import Document


DATA_FILES = [
    r"d:\Client Projects\jobiho\src\data\destinations.ts",
    r"d:\Client Projects\jobiho\src\data\destinations-americas.ts",
    r"d:\Client Projects\jobiho\src\data\destinations-meea.ts",
    r"d:\Client Projects\jobiho\src\data\destinations-east-asia.ts",
    r"d:\Client Projects\jobiho\src\data\destinations-south-southeast-asia.ts",
]

OUTPUT_PATH = r"d:\Client Projects\jobiho\Tours_and_Packages_Master_Final.docx"


def extract_destinations_from_text(text: str):
    pattern = re.compile(
        r'name:\s*"(?P<name>[^"]+)"'
        r'.*?country:\s*"(?P<country>[^"]+)"'
        r'.*?region:\s*"(?P<region>[^"]+)"'
        r'.*?itinerary:\s*\[(?P<itinerary>.*?)\]\s*,\s*thingsToDo:',
        re.DOTALL,
    )
    for m in pattern.finditer(text):
        itinerary_entries = re.findall(
            r'\{\s*day:\s*(\d+),\s*title:\s*"([^"]+)"',
            m.group("itinerary"),
            re.DOTALL,
        )
        if not itinerary_entries:
            continue
        yield {
            "place": m.group("name").strip(),
            "country": title_case_country(m.group("country").strip()),
            "continent": region_to_continent(
                m.group("region").strip(), m.group("country").strip()
            ),
            "day_activities": [(int(day), title) for day, title in itinerary_entries],
        }


def title_case_country(country: str) -> str:
    if country.upper() == "UAE":
        return "UAE"
    if country == "United States":
        return "United States of America"
    if country == "USA":
        return "United States of America"
    return country


def americas_subregion(country: str) -> str:
    north = {
        "United States",
        "United States of America",
        "USA",
        "Canada",
        "Mexico",
    }
    south = {
        "Peru",
        "Brazil",
        "Argentina",
        "Chile",
        "Colombia",
        "Ecuador",
        "Bolivia",
        "Uruguay",
        "Paraguay",
        "Venezuela",
    }
    if country in north:
        return "North America"
    if country in south:
        return "South America"
    return "Americas"


def region_to_continent(region: str, country: str) -> str:
    r = region.strip().lower()
    if r == "americas":
        return americas_subregion(country)
    if r in {"europe", "asia", "africa", "oceania"}:
        return region
    if r in {"middle east", "meea", "middle east & africa", "middle east and africa"}:
        # Split by country for cleaner hierarchy.
        middle_east = {
            "UAE",
            "United Arab Emirates",
            "Jordan",
            "Oman",
            "Qatar",
            "Saudi Arabia",
            "Turkey",
        }
        if country in middle_east:
            return "Middle East"
        return "Africa"
    return region


def parse_destinations():
    items = []
    for path in DATA_FILES:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
        items.extend(extract_destinations_from_text(text))
    # stable sorted output
    items.sort(key=lambda x: (x["continent"], x["country"], x["place"]))
    return items


def build_doc(items):
    doc = Document()
    doc.add_heading("Tours and Packages Master List", 0)
    doc.add_paragraph(
        "Format: World -> Region -> Nation -> Place -> Day-based activity/place name"
    )
    doc.add_paragraph("Photos needed per place: 5 random aesthetic location photos")

    hierarchy = defaultdict(lambda: defaultdict(list))
    for item in items:
        hierarchy[item["continent"]][item["country"]].append(item)

    region_labels = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    preferred_region_order = [
        "North America",
        "South America",
        "Europe",
        "Asia",
        "Middle East",
        "Africa",
        "Oceania",
        "Americas",
    ]
    all_regions = sorted(hierarchy.keys())
    ordered_regions = [r for r in preferred_region_order if r in hierarchy] + [
        r for r in all_regions if r not in preferred_region_order
    ]
    for ridx, region in enumerate(ordered_regions):
        label = region_labels[ridx] if ridx < len(region_labels) else f"R{ridx+1}"
        doc.add_paragraph(f"{label}) World -> {region}")
        countries = sorted(hierarchy[region].keys())
        for cidx, country in enumerate(countries, start=1):
            country_package_count = len(hierarchy[region][country])
            doc.add_paragraph(
                f"  {to_roman(cidx)}) World -> {region} -> {country} ({country_package_count} packages/places)"
            )
            places = sorted(hierarchy[region][country], key=lambda x: x["place"])
            for pidx, place in enumerate(places, start=1):
                doc.add_paragraph(
                    f"    {pidx}) {place['place']} (5 random aesthetic location photos)"
                )
                for day, title in sorted(place["day_activities"], key=lambda x: x[0]):
                    doc.add_paragraph(f"       *) Day {day}: {title}")
                doc.add_paragraph("")

    doc.save(OUTPUT_PATH)


def to_roman(num: int) -> str:
    vals = [
        (1000, "m"),
        (900, "cm"),
        (500, "d"),
        (400, "cd"),
        (100, "c"),
        (90, "xc"),
        (50, "l"),
        (40, "xl"),
        (10, "x"),
        (9, "ix"),
        (5, "v"),
        (4, "iv"),
        (1, "i"),
    ]
    out = []
    n = num
    for v, sym in vals:
        while n >= v:
            out.append(sym)
            n -= v
    return "".join(out)


def main():
    items = parse_destinations()
    build_doc(items)
    print(f"Created: {OUTPUT_PATH}")
    print(f"Total places: {len(items)}")


if __name__ == "__main__":
    main()
